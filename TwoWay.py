import streamlit as st
import pandas as pd
import numpy as np
import statsmodels.api as sm
from statsmodels.formula.api import ols
import matplotlib.pyplot as plt
import seaborn as sns
from statsmodels.stats.multicomp import pairwise_tukeyhsd, MultiComparison
import plotly.express as px
import plotly.graph_objects as go
from scipy import stats
import datetime
import io
from docx import Document

# Add these functions before your main code

def format_p_value(p_value):
    """Format p-value with appropriate precision based on magnitude"""
    return f"{p_value:.4f}" if p_value >= 0.0001 else f"{p_value:.4e}"

def perform_tukey_hsd(data, factor, dependent_var, alpha):
    """Perform Tukey HSD test"""
    tukey = pairwise_tukeyhsd(endog=data[dependent_var], groups=data[factor], alpha=alpha)
    results = pd.DataFrame(
        data=tukey._results_table.data[1:],
        columns=tukey._results_table.data[0]
    )
    return results, "Tukey HSD"

def perform_bonferroni(data, factor, dependent_var, alpha):
    """Perform Bonferroni correction"""
    mc = MultiComparison(data[dependent_var], data[factor])
    result_obj = mc.allpairtest(stats.ttest_ind, method='bonf')
    
    # Convert to DataFrame similar to Tukey output
    df = pd.DataFrame(columns=['group1', 'group2', 'meandiff', 'p-adj', 'lower', 'upper', 'reject'])
    
    # Extract data from the MultiComparison result object
    raw_data = result_obj[0].data
    pvalues = result_obj[0].pvalues
    reject = result_obj[0].reject
    
    # Get the group comparison pairs
    rows = []
    for i, (group1, group2) in enumerate(mc.pairindices):
        g1 = mc.groupsunique[group1]
        g2 = mc.groupsunique[group2]
        
        # Calculate mean difference
        mean1 = data[data[factor] == g1][dependent_var].mean()
        mean2 = data[data[factor] == g2][dependent_var].mean()
        meandiff = mean1 - mean2
        
        # Calculate confidence intervals
        n1 = len(data[data[factor] == g1])
        n2 = len(data[data[factor] == g2])
        
        # Calculate pooled standard deviation
        var1 = data[data[factor] == g1][dependent_var].var()
        var2 = data[data[factor] == g2][dependent_var].var()
        df_pooled = n1 + n2 - 2
        
        pooled_sd = np.sqrt(((n1-1) * var1 + (n2-1) * var2) / df_pooled)
        
        # Standard error
        se = pooled_sd * np.sqrt(1/n1 + 1/n2)
        
        # Critical value adjusted for Bonferroni
        t_crit = stats.t.ppf(1 - alpha/(2*len(mc.pairindices)), df_pooled)
        
        # Confidence interval
        lower = meandiff - t_crit * se
        upper = meandiff + t_crit * se
        
        # Add to results
        rows.append({
            'group1': g1, 
            'group2': g2, 
            'meandiff': meandiff,
            'p-adj': pvalues[i] if i < len(pvalues) else np.nan,
            'lower': lower,
            'upper': upper,
            'reject': reject[i] if i < len(reject) else np.nan
        })
    
    # Create DataFrame from rows
    if rows:
        df = pd.DataFrame(rows)
    
    return df, "Bonferroni"

def perform_scheffe(data, factor, dependent_var, alpha):
    """Perform Scheffé test"""
    # Fit ANOVA model
    formula = f"{dependent_var} ~ C({factor})"
    model = ols(formula, data=data).fit()
    anova_results = sm.stats.anova_lm(model, typ=2)
    
    # Get MSE and degrees of freedom
    mse = anova_results.loc['Residual', 'sum_sq'] / anova_results.loc['Residual', 'df']
    df_error = anova_results.loc['Residual', 'df']
    
    # Get all group combinations
    groups = data[factor].unique()
        n2 = len(data[data[factor] == g2])
        mean1 = data[data[factor] == g1][dependent_var].mean()
        mean2 = data[data[factor] == g2][dependent_var].mean()
        meandiff = mean1 - mean2
        
        # Scheffé test statistic
        se = np.sqrt(mse * (1/n1 + 1/n2))
        scheffe_stat = meandiff**2 / (mse * (1/n1 + 1/n2) * (k-1))
        p_val = 1 - stats.f.cdf(scheffe_stat, k-1, df_error)
        
        # Confidence interval
        ci_margin = np.sqrt(f_crit) * se
        lower = meandiff - ci_margin
        upper = meandiff + ci_margin
        
        # Reject null if F statistic > F critical
        reject = scheffe_stat > f_crit
        
        results.append([g1, g2, meandiff, p_val, lower, upper, reject])
    
    df = pd.DataFrame(results, columns=['group1', 'group2', 'meandiff', 'p-adj', 'lower', 'upper', 'reject'])
    return df, "Scheffé"

def perform_games_howell(data, factor, dependent_var, alpha):
    """Perform Games-Howell test for unequal variances"""
    groups = data[factor].unique()
    combinations = [(g1, g2) for i, g1 in enumerate(groups) for g2 in groups[i+1:]]
    
    results = []
    for g1, g2 in combinations:
        # Get group data
        group1 = data[data[factor] == g1][dependent_var]
        group2 = data[data[factor] == g2][dependent_var]
        
        # Group stats
        n1, n2 = len(group1), len(group2)
        mean1, mean2 = group1.mean(), group2.mean()
        var1, var2 = group1.var(), group2.var()
        
        # Mean difference
        meandiff = mean1 - mean2
        
        # Welch-Satterthwaite degrees of freedom
        df_num = (var1/n1 + var2/n2)**2
        df_denom = (var1/n1)**2/(n1-1) + (var2/n2)**2/(n2-1)
        df = df_num / df_denom if df_denom > 0 else np.inf
        
        # Standard error and t-statistic
        se = np.sqrt(var1/n1 + var2/n2)
        t_stat = meandiff / se
        
        # p-value and confidence interval
        p_val = 2 * (1 - stats.t.cdf(abs(t_stat), df))
        q_crit = stats.studentized_range.ppf(1-alpha, len(groups), df) / np.sqrt(2)
        ci_margin = q_crit * se
        lower = meandiff - ci_margin
        upper = meandiff + ci_margin
        
        # Reject null if |t| > critical value
        reject = abs(t_stat) > q_crit
        
        results.append([g1, g2, meandiff, p_val, lower, upper, reject])
    
    df = pd.DataFrame(results, columns=['group1', 'group2', 'meandiff', 'p-adj', 'lower', 'upper', 'reject'])
    return df, "Games-Howell"

def create_posthoc_analysis(data, factor, dependent_var, alpha, method='tukey'):
    """Perform post-hoc analysis for a factor using the specified method"""
    # Select the appropriate post-hoc test method
    if method == 'bonferroni':
        results, method_name = perform_bonferroni(data, factor, dependent_var, alpha)
    elif method == 'scheffe':
        results, method_name = perform_scheffe(data, factor, dependent_var, alpha)
    elif method == 'games_howell':
        results, method_name = perform_games_howell(data, factor, dependent_var, alpha)
    else:  # Default to Tukey HSD
        results, method_name = perform_tukey_hsd(data, factor, dependent_var, alpha)
    
    st.write(f"#### {method_name} Post-hoc Test Results")
    st.write(results)
    
    # Create visualizations
    st.write(f"### Visualisasi Efek {factor}")
    
    # Boxplot
    fig = px.box(
        data, 
        x=factor, 
        y=dependent_var,
        title=f"Boxplot {dependent_var} berdasarkan {factor}",
        color=factor
    )
    st.plotly_chart(fig, use_container_width=True)
    
    # Bar plot with error bars
    summary_stats = data.groupby(factor, observed=True)[dependent_var].agg(['mean', 'std']).reset_index()
    
    fig = px.bar(
        summary_stats,
        x=factor,
        y='mean',
        error_y='std',
        title=f"Rata-rata {dependent_var} berdasarkan {factor} (dengan standar deviasi)",
        color=factor
    )
    st.plotly_chart(fig, use_container_width=True)
    
    return results

# Keep the original create_tukey_analysis function for backward compatibility
def create_tukey_analysis(data, factor, dependent_var, alpha):
    """Perform Tukey HSD analysis for a factor (Legacy function)"""
    return create_posthoc_analysis(data, factor, dependent_var, alpha, method='tukey')

def create_docx_report(data, dependent_var, factor1, factor2, formatted_anova, effects_summary, p_values, eta_squared, alpha, model, conclusions):
    """Generate a Word document report for Two-Way ANOVA analysis."""
    doc = Document()
    
    # Add title and date
    doc.add_heading('Two-Way ANOVA Analysis Report', 0)
    doc.add_paragraph(f'Generated on: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Analysis overview
    doc.add_heading('Analysis Overview', 1)
    doc.add_paragraph(f'Dependent Variable: {dependent_var}')
    doc.add_paragraph(f'Factor 1: {factor1}')
    doc.add_paragraph(f'Factor 2: {factor2}')
    doc.add_paragraph(f'Significance level (α): {alpha}')
    doc.add_paragraph(f'Sample size: {len(data)}')
    
    # ANOVA Table
    doc.add_heading('ANOVA Results', 1)
    
    # Add ANOVA table
    table = doc.add_table(rows=len(formatted_anova)+1, cols=len(formatted_anova.columns)+1)
    table.style = 'Table Grid'
    
    # Add headers
    headers = ['Factor'] + list(formatted_anova.columns)
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    
    # Add data
    for i, (idx, row) in enumerate(formatted_anova.iterrows(), 1):
        table.cell(i, 0).text = idx
        for j, val in enumerate(row, 1):
            table.cell(i, j).text = str(val)
    
    # Effect Sizes
    doc.add_heading('Effect Sizes', 1)
    
    # Add effect size table
    table = doc.add_table(rows=len(effects_summary)+1, cols=len(effects_summary.columns))
    table.style = 'Table Grid'
    
    # Add headers
    for i, header in enumerate(effects_summary.columns):
        table.cell(0, i).text = header
    
    # Add data
    for i, row in enumerate(effects_summary.values, 1):
        for j, val in enumerate(row):
            table.cell(i, j).text = str(val)
    
    # Conclusions
    doc.add_heading('Conclusions', 1)
    for conclusion in conclusions:
        doc.add_paragraph(conclusion, style='List Bullet')
    
    doc.add_paragraph(f"Model explains {model.rsquared:.2%} of the variation in {dependent_var}.")
    
    # Add interaction note if significant
    if p_values[2] < alpha:
        doc.add_paragraph("Interaction Effect Note:", style='Intense Quote')
        doc.add_paragraph("When an interaction is significant, the effect of one factor depends on the level of the other factor. In this case, the simple main effects should be interpreted with caution.")
    
    # Add footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Generated using Two-Way ANOVA Analysis Tool | © {datetime.datetime.now().year} Galuh Adi Insani"
    
    # Save to BytesIO object
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    
    return docx_io

def create_html_report(dependent_var, factor1, factor2, formatted_anova, effects_summary, p_values, eta_squared, alpha, model, conclusions):
    """Generate an HTML report for Two-Way ANOVA analysis that can be opened in Word."""
    
    # Convert dataframes to HTML tables with bootstrap styling
    anova_table_html = formatted_anova.to_html(index=True, classes="table table-striped")
    effects_table_html = effects_summary.to_html(index=False, classes="table table-striped")
    
    # Format conclusions as list items
    conclusions_html = "".join([f"<li>{c}</li>" for c in conclusions])
    
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Two-Way ANOVA Analysis Report</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; }}
            h1 {{ color: #2C3E50; }}
            h2 {{ color: #34495E; margin-top: 30px; }}
            table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
            table, th, td {{ border: 1px solid #ddd; }}
            th, td {{ padding: 12px; text-align: left; }}
            th {{ background-color: #f2f2f2; }}
            .footer {{ margin-top: 40px; text-align: center; font-size: 12px; color: #7F8C8D; }}
        </style>
    </head>
    <body>
        <h1>Two-Way ANOVA Analysis Report</h1>
        <p>Generated on: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
        
        <h2>Analysis Overview</h2>
        <p><strong>Dependent Variable:</strong> {dependent_var}</p>
        <p><strong>Factor 1:</strong> {factor1}</p>
        <p><strong>Factor 2:</strong> {factor2}</p>
        <p><strong>Significance level (α):</strong> {alpha}</p>
        
        <h2>ANOVA Results</h2>
        {anova_table_html}
        
        <h2>Effect Sizes</h2>
        {effects_table_html}
        
        <h2>Conclusions</h2>
        <ul>
            {conclusions_html}
        </ul>
        <p><strong>Model explains {model.rsquared:.2%} of the variation in {dependent_var}.</strong></p>
        
        <div class="footer">
            <p>Generated using Two-Way ANOVA Analysis Tool | © {datetime.datetime.now().year} Galuh Adi Insani</p>
        </div>
    </body>
    </html>
    """
    
    return html

def validate_and_clean_data(data, dependent_var, factor1, factor2):
    """
    Validate and clean data for ANOVA analysis by removing NaN/Inf values
    and checking for appropriate group sizes.
    
    Returns:
        tuple: (clean_data, validation_message)
    """
    validation_message = []
    
    # Check for NaN/Inf in the dependent variable
    invalid_values = data[dependent_var].isna() | np.isinf(data[dependent_var])
    if invalid_values.any():
        rows_with_invalid = invalid_values.sum()
        validation_message.append(f"Removed {rows_with_invalid} rows with NaN or Inf values in {dependent_var}.")
        data = data[~invalid_values]
    
    # Check for NaN in the factors
    for factor in [factor1, factor2]:
        invalid_factor = data[factor].isna()
        if invalid_factor.any():
            rows_with_invalid = invalid_factor.sum()
            validation_message.append(f"Removed {rows_with_invalid} rows with NaN values in {factor}.")
            data = data[~invalid_factor]
    
    # Check minimum group size
    group_sizes = data.groupby([factor1, factor2]).size()
    min_group_size = group_sizes.min()
    if min_group_size < 3:
        validation_message.append(f"Warning: Some groups have fewer than 3 observations (minimum: {min_group_size}).")
        small_groups = group_sizes[group_sizes < 3].reset_index()
        for _, row in small_groups.iterrows():
            validation_message.append(f"  - Group {factor1}={row[factor1]}, {factor2}={row[factor2]} has only {row[0]} observation(s).")
    
    # Check if any groups have zero variance
    zero_var_groups = []
    for name, group in data.groupby([factor1, factor2]):
        if group[dependent_var].var() == 0:
            zero_var_groups.append(name)
    
    if zero_var_groups:
        validation_message.append(f"Warning: Some groups have zero variance in {dependent_var}:")
        for group in zero_var_groups:
            validation_message.append(f"  - Group {factor1}={group[0]}, {factor2}={group[1]}")
    
    return data, validation_message

# Set page configuration
st.set_page_config(
    page_title="Analisis Two-Way ANOVA",
    page_icon="📊",
    layout="wide"
)

st.title('📊 Analisis Two-Way ANOVA')
st.markdown("""
    Aplikasi ini membantu Anda melakukan analisis Two-Way ANOVA pada data Anda.
    Upload file CSV atau Excel, pilih variabel yang ingin dianalisis, dan dapatkan hasil lengkap.
""")

# Hide default Streamlit elements
hide_st_style = """
        <style>
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}
        </style>
        """
st.markdown(hide_st_style, unsafe_allow_html=True)

# Sidebar for additional options
with st.sidebar:
    st.header("Informasi")
    st.info("""
        **Two-Way ANOVA** digunakan untuk menilai pengaruh dari dua variabel kategori 
        independen (faktor) terhadap satu variabel dependen kontinu.
        
        Aplikasi ini akan menampilkan:
        - Tabel ANOVA
        - Ukuran efek
        - Tes post-hoc
        - Visualisasi data
    """)
    
    st.header("Pengaturan")
    alpha = st.slider("Tingkat alpha", 0.01, 0.10, 0.05, 0.01)
    show_assumptions = st.checkbox("Periksa asumsi ANOVA", True)
    
    # Add selection for post-hoc test method
    posthoc_method = st.selectbox(
        "Metode Post-hoc Test",
        ['tukey', 'bonferroni', 'scheffe', 'games_howell'],
        format_func=lambda x: {
            'tukey': 'Tukey HSD (Equal variances)',
            'bonferroni': 'Bonferroni (Conservative)',
            'scheffe': 'Scheffé (Flexible contrasts)',
            'games_howell': 'Games-Howell (Unequal variances)'
        }[x],
        help="""
        - Tukey HSD: Standard post-hoc test for equal variances
        - Bonferroni: More conservative test (controls family-wise error)
        - Scheffé: Allows for complex comparisons between groups
        - Games-Howell: Better for unequal variances between groups
        """
    )

# Upload file CSV atau Excel
uploaded_file = st.file_uploader("Unggah file CSV atau Excel Anda", type=["csv", "xlsx", "xls"])

if uploaded_file is not None:
    try:
        # Determine file type and read accordingly
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension == 'csv':
            data = pd.read_csv(uploaded_file)
        elif file_extension in ['xlsx', 'xls']:
            data = pd.read_excel(uploaded_file)
        
        # Add this within your data processing section after reading the data

        # Detect if dataset is large and offer sampling
        if data.shape[0] > 10000:
            st.warning(f"Dataset Anda cukup besar ({data.shape[0]} baris). Pertimbangkan untuk menggunakan sampel untuk analisis yang lebih cepat.")
            use_sample = st.checkbox("Gunakan sampel data untuk analisis", True)
            if use_sample:
                sample_size = st.slider("Ukuran sampel", min_value=1000, max_value=min(10000, data.shape[0]), value=5000, step=1000)
                data = data.sample(n=sample_size, random_state=42)
                st.success(f"Menggunakan {sample_size} sampel data untuk analisis.")
        
        # Tampilkan informasi data
        col1, col2 = st.columns([2, 1])
        with col1:
            st.write("### Pratinjau Data:")
            st.write(data.head())
        
        with col2:
            st.write("### Informasi Data:")
            st.write(f"Jumlah baris: {data.shape[0]}")
            st.write(f"Jumlah kolom: {data.shape[1]}")
            
            # Deteksi tipe data untuk membantu pemilihan kolom
            numeric_cols = data.select_dtypes(include=['number']).columns.tolist()
            categorical_cols = data.select_dtypes(include=['object', 'category']).columns.tolist()
            
            if len(categorical_cols) < 2:
                st.warning("Data Anda memiliki kurang dari 2 kolom kategorikal. Pertimbangkan untuk mengubah tipe data kolom yang sesuai.")

        # Tambahkan expander untuk informasi lengkap data
        with st.expander("Informasi lengkap data"):
            buffer = pd.DataFrame({
                'Kolom': data.columns,
                'Tipe Data': [str(dtype) for dtype in data.dtypes],  # Konversi dtypes ke string
                'Nilai Unik': [data[col].nunique() for col in data.columns],
                'Missing Values': data.isnull().sum().values,
            })
            st.write(buffer)

        # Pilih kolom untuk analisis dengan pengelompokan yang lebih jelas
        st.write("### Pemilihan Variabel")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            # Sarankan kolom numerik untuk variabel dependen
            dependent_var = st.selectbox(
                "Pilih variabel dependen (numerik)", 
                options=numeric_cols if numeric_cols else data.columns.tolist(),
                help="Variabel dependen harus bertipe numerik"
            )
        
        with col2:
            # Sarankan kolom kategorikal untuk faktor
            factor1 = st.selectbox(
                "Pilih faktor pertama (kategorikal)", 
                options=categorical_cols if categorical_cols else data.columns.tolist(),
                help="Faktor pertama sebaiknya bertipe kategorikal"
            )
        
        with col3:
            # Sarankan kolom kategorikal untuk faktor kedua
            factor2 = st.selectbox(
                "Pilih faktor kedua (kategorikal)", 
                options=[col for col in (categorical_cols if categorical_cols else data.columns.tolist()) if col != factor1],
                help="Faktor kedua sebaiknya bertipe kategorikal dan berbeda dengan faktor pertama"
            )
        
        # Pastikan kolom terpilih valid
        if factor1 == factor2:
            st.error("Faktor pertama dan kedua tidak boleh sama.")
        
        # Tombol untuk menjalankan analisis
        run_button = st.button("Jalankan Two-Way ANOVA", type="primary")
        
        # Add this after loading the data but before analysis

        # Check for appropriate variable types
        if run_button:
            with st.spinner('Validasi data...'):
                # Check minimum group size
                min_group_size = data.groupby([factor1, factor2]).size().min()
                if min_group_size < 3:
                    st.warning(f"Beberapa kelompok memiliki ukuran sampel kurang dari 3 (minimum: {min_group_size}). Hasil ANOVA mungkin tidak reliable.")
                
                # Check if factors have adequate levels
                if data[factor1].nunique() < 2:
                    st.error(f"Faktor {factor1} harus memiliki minimal 2 level. Saat ini hanya memiliki {data[factor1].nunique()} level.")
                    st.stop()
                    
                if data[factor2].nunique() < 2:
                    st.error(f"Faktor {factor2} harus memiliki minimal 2 level. Saat ini hanya memiliki {data[factor2].nunique()} level.")
                    st.stop()
        
        if run_button:
            with st.spinner('Sedang memproses...'):
                # Validasi data
                if data[dependent_var].dtype.kind not in 'if':  # Cek apakah numerik (integer atau float)
                    st.error(f"Variabel dependen '{dependent_var}' harus bertipe numerik. Tipe data saat ini: {data[dependent_var].dtype}")
                    st.stop()
                
                # Konversi faktor ke kategorikal jika belum
                data[factor1] = data[factor1].astype('category')
                data[factor2] = data[factor2].astype('category')
                
                # Cek data kosong
                if data[[dependent_var, factor1, factor2]].isnull().any().any():
                    st.warning("Data Anda mengandung nilai kosong (NA/null). Baris dengan nilai kosong akan dihapus.")
                    data = data.dropna(subset=[dependent_var, factor1, factor2])
                
                # Membuat dataframe sementara dengan kolom yang diganti nama untuk menghindari masalah keyword
                temp_data = data.copy()
                column_mapping = {
                    dependent_var: 'dependent_var',
                    factor1: 'factor1',
                    factor2: 'factor2'
                }
                temp_data = temp_data.rename(columns=column_mapping)
                
                # Cek asumsi ANOVA jika diminta
                if show_assumptions:
                    st.write("## Uji Asumsi ANOVA")
                    
                    # Uji normalitas (Shapiro-Wilk)
                    st.write("### 1. Uji Normalitas")
                    
                    groups = []
                    for f1 in data[factor1].unique():
                        for f2 in data[factor2].unique():
                            subset = data[(data[factor1] == f1) & (data[factor2] == f2)][dependent_var]
                            if not subset.empty and len(subset) >= 3:  # Shapiro-Wilk membutuhkan minimal 3 nilai
                                try:
                                    stat, p = stats.shapiro(subset)
                                    groups.append((f1, f2, stat, p, p > alpha))
                                except Exception as e:
                                    st.warning(f"Could not perform Shapiro-Wilk test on group {f1}, {f2}: {str(e)}")
                                    # Skip this group but continue with others
                    
                    if groups:
                        assumption_df = pd.DataFrame(groups, columns=['Faktor 1', 'Faktor 2', 'Statistik', 'P-value', 'Normal'])
                        st.write(assumption_df)
                        
                        if not all(assumption_df['Normal']):
                            st.warning("Beberapa kelompok tidak memenuhi asumsi normalitas. Pertimbangkan untuk menggunakan tes non-parametrik.")
                        else:
                            st.success("Asumsi normalitas terpenuhi untuk semua kelompok.")
                    else:
                        st.warning("Tidak dapat menguji normalitas - beberapa kelompok memiliki terlalu sedikit data.")
                    
                    # Uji homogenitas varians (Levene)
                    st.write("### 2. Uji Homogenitas Varians")
                    try:
                        # Buat kolom gabungan untuk Levene test
                        temp_data['group'] = temp_data['factor1'].astype(str) + "_" + temp_data['factor2'].astype(str)
                        
                        # Pendekatan yang lebih aman: filter data terlebih dahulu
                        levene_data = []
                        levene_groups = []
                        
                        for group in temp_data['group'].unique():
                            group_data = temp_data[temp_data['group'] == group]['dependent_var']
                            # Filter: minimal 3 sampel dan variansi yang cukup
                            if len(group_data) >= 3 and group_data.var() > 0.0001:
                                levene_data.append(group_data.values)
                                levene_groups.append(group)
                        
                        if len(levene_data) >= 2:  # Levene test memerlukan minimal 2 grup
                            # Tangani peringatan dengan context manager
                            import warnings
                            with warnings.catch_warnings():
                                warnings.filterwarnings("ignore", category=RuntimeWarning)
                                stat, p = stats.levene(*levene_data)
                            
                            st.write(f"Statistik Levene: {stat:.4f}")
                            st.write(f"P-value: {p:.4f}")
                            
                            if p > alpha:
                                st.success(f"Asumsi homogenitas varians terpenuhi (p = {p:.4f} > {alpha}).")
                            else:
                                st.warning(f"Asumsi homogenitas varians tidak terpenuhi (p = {p:.4f} < {alpha}). Pertimbangkan untuk menggunakan koreksi.")
                        else:
                            st.warning("Tidak dapat menguji homogenitas varians - tidak cukup kelompok dengan variansi signifikan.")
                    except Exception as e:
                        st.warning(f"Tidak dapat melakukan uji homogenitas varians: {str(e)}")
                
                # Jalankan Two-Way ANOVA
                st.write("## Hasil Two-Way ANOVA")
                
                # Gunakan kolom yang sudah diganti nama dalam formula
                formula = "dependent_var ~ C(factor1) + C(factor2) + C(factor1):C(factor2)"
                try:
                    model = ols(formula, data=temp_data).fit()
                    anova_table = sm.stats.anova_lm(model, typ=2)
                except Exception as e:
                    st.error(f"ANOVA calculation failed: {str(e)}")
                    st.info("This may be due to perfect multicollinearity, empty groups, or non-finite values in the data.")
                    st.stop()
                
                # Format tabel ANOVA
                anova_display = anova_table.copy()
                anova_display.index = [
                    f"{factor1}" if idx == "C(factor1)" else
                    f"{factor2}" if idx == "C(factor2)" else
                    f"Interaksi ({factor1}:{factor2})" if idx == "C(factor1):C(factor2)" else
                    "Residual"
                    for idx in anova_display.index
                ]
                
                # Tampilkan tabel ANOVA dengan format yang lebih baik
                st.write("### Tabel ANOVA:")
                
                # Format angka dalam tabel ANOVA untuk tampilan yang lebih baik
                formatted_anova = anova_display.copy()
                
                # Perbaikan 1: Periksa kolom yang ada sebelum memformat
                for col in formatted_anova.columns:
                    if col in ['sum_sq', 'df', 'F']:
                        formatted_anova[col] = formatted_anova[col].apply(
                            lambda x: f"{x:.4f}" if isinstance(x, (float, np.floating)) else str(x)
                        )
                
                # Hitung mean_sq jika tidak ada dalam tabel
                if 'mean_sq' not in formatted_anova.columns:
                    # Tambahkan kolom mean_sq yang dihitung dari sum_sq/df
                    formatted_anova['mean_sq'] = formatted_anova['sum_sq'].astype(float) / formatted_anova['df'].astype(float)
                    formatted_anova = formatted_anova[['sum_sq', 'df', 'mean_sq', 'F', 'PR(>F)']]  # Reorder columns
                
                # Format kolom mean_sq setelah perhitungan
                formatted_anova['mean_sq'] = formatted_anova['mean_sq'].apply(
                    lambda x: f"{x:.4f}" if isinstance(x, (float, np.floating)) else str(x)
                )
                
                # Format p-value dengan notasi ilmiah untuk nilai yang sangat kecil
                formatted_anova['PR(>F)'] = formatted_anova['PR(>F)'].apply(lambda x: f"{x:.4f}" if x >= 0.0001 else f"{x:.4e}")
                
                st.write(formatted_anova)
                
                # Tambahkan unduhan tabel ANOVA
                csv = formatted_anova.to_csv(index=True)
                st.download_button(
                    label="Unduh Tabel ANOVA sebagai CSV",
                    data=csv,
                    file_name="anova_results.csv",
                    mime="text/csv"
                )
                
                # Tampilkan hasil dengan nama kolom asli
                st.write("### Ukuran Efek")
                p_values = anova_table['PR(>F)']
                
                # Hitung ukuran efek (Eta-squared dan Partial Eta-squared)
                ss_total = anova_table['sum_sq'].sum()
                eta_squared = anova_table['sum_sq'] / ss_total
                
                # Hitung partial eta-squared
                partial_eta_squared = anova_table['sum_sq'] / (anova_table['sum_sq'] + anova_table['sum_sq'].iloc[-1])
                partial_eta_squared = partial_eta_squared[:-1]  # Hapus residu
                
                # Buat tabel hasil efek
                effects_summary = pd.DataFrame({
                    'Faktor': [f"{factor1}", f"{factor2}", f"Interaksi ({factor1}:{factor2})"],
                    'p-value': p_values[:-1].values,
                    'Signifikan': p_values[:-1].values < alpha,
                    'Eta²': eta_squared[:-1].values,
                    'Partial Eta²': partial_eta_squared.values,
                    'Efek': [
                        "Besar" if eta_squared.iloc[0] >= 0.14 else "Sedang" if eta_squared.iloc[0] >= 0.06 else "Kecil",
                        "Besar" if eta_squared.iloc[1] >= 0.14 else "Sedang" if eta_squared.iloc[1] >= 0.06 else "Kecil",
                        "Besar" if eta_squared.iloc[2] >= 0.14 else "Sedang" if eta_squared.iloc[2] >= 0.06 else "Kecil"
                    ]
                })
                
                # Format angka dalam tabel efek
                effects_summary['p-value'] = effects_summary['p-value'].apply(lambda x: f"{x:.4f}" if x >= 0.0001 else f"{x:.4e}")
                effects_summary['Eta²'] = effects_summary['Eta²'].apply(
                    lambda x: f"{x:.4f}" if isinstance(x, (float, np.floating)) else str(x)
                )
                effects_summary['Partial Eta²'] = effects_summary['Partial Eta²'].apply(
                    lambda x: f"{x:.4f}" if isinstance(x, (float, np.floating)) else str(x)
                )
                
                st.write(effects_summary)
                
                # Tambahkan visualisasi ringkasan hasil
                st.write("### Visualisasi Ukuran Efek")
                
                fig = px.bar(
                    effects_summary, 
                    x='Faktor', 
                    y='Eta²',
                    color='Signifikan',
                    text='Eta²',
                    title='Ukuran Efek (Eta²) untuk Setiap Faktor',
                    color_discrete_map={True: 'green', False: 'gray'}
                )
                fig.update_layout(height=500)
                st.plotly_chart(fig, use_container_width=True)
                
                # Tambahkan ringkasan model dengan tab
                st.write("### Ringkasan Model")
                tabs = st.tabs(["Koefisien Model", "Statistik Model", "Ringkasan Lengkap"])
                
                with tabs[0]:
                    coef_df = pd.DataFrame({
                        'Koefisien': model.params,
                        'Std Error': model.bse,
                        't-value': model.tvalues,
                        'p-value': model.pvalues
                    })
                    st.write(coef_df)
                
                with tabs[1]:
                    stats_df = pd.DataFrame({
                        'Metrik': ['R²', 'R² Disesuaikan', 'F-statistic', 'p-value', 'Log-Likelihood', 'AIC', 'BIC'],
                        'Nilai': [
                            f"{model.rsquared:.4f}",
                            f"{model.rsquared_adj:.4f}",
                            f"{model.fvalue:.4f}",
                            f"{model.f_pvalue:.4e}" if model.f_pvalue < 0.0001 else f"{model.f_pvalue:.4f}",
                            f"{model.llf:.4f}",
                            f"{model.aic:.4f}",
                            f"{model.bic:.4f}"
                        ]
                    })
                    st.write(stats_df)
                
                with tabs[2]:
                    st.text(str(model.summary()))
                
                # Tambahkan interpretasi detail
                st.write("### Interpretasi Detail")
                
                st.write("**Interpretasi Ukuran Efek:**")
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.info("η² < 0.06: Efek kecil")
                with col2:
                    st.info("0.06 ≤ η² < 0.14: Efek sedang")
                with col3:
                    st.info("η² ≥ 0.14: Efek besar")
                
                # Uji post-hoc untuk efek utama yang signifikan
                st.write("## Uji Post-hoc")

                with st.expander("Informasi tentang uji post-hoc"):
                    st.markdown("""
                    ### Panduan Memilih Uji Post-hoc

                    Berbagai uji post-hoc memiliki karakteristik berbeda:

                    | Uji Post-hoc | Kelebihan | Kekurangan | Disarankan Untuk |
                    |--------------|-----------|------------|------------------|
                    | **Tukey HSD** | Kekuatan statistik yang baik | Mengasumsikan variansi sama | Ukuran sampel sama dan variansi homogen |
                    | **Bonferroni** | Sangat ketat dalam mengontrol Type I error | Bisa terlalu konservatif, kurang powerful | Dataset dengan banyak perbandingan |
                    | **Scheffé** | Cocok untuk semua jenis perbandingan dan kontras | Kurang powerful dari Tukey | Pengujian kompleks dengan berbagai kontras |
                    | **Games-Howell** | Tidak mengasumsikan variansi sama | Perhitungan lebih kompleks | Data dengan variansi tidak homogen |

                    **Situasi khusus:**
                    - Jika asumsi homogenitas variansi dilanggar, gunakan Games-Howell
                    - Jika ada banyak perbandingan yang akan dilakukan, Bonferroni dapat membantu mengurangi false positives
                    - Jika ukuran grup sangat tidak seimbang, Games-Howell juga merupakan pilihan yang baik
                    """)

                # Jalankan post-hoc test yang dipilih untuk efek utama yang signifikan
                if p_values.iloc[0] < alpha:
                    st.write(f"### {factor1}")
                    create_posthoc_analysis(data, factor1, dependent_var, alpha, method=posthoc_method)
                else:
                    st.info(f"Faktor {factor1} tidak memiliki efek yang signifikan (p = {p_values.iloc[0]:.4f} > {alpha}).")
                
                if p_values.iloc[1] < alpha:
                    st.write(f"### {factor2}")
                    create_posthoc_analysis(data, factor2, dependent_var, alpha, method=posthoc_method)
                else:
                    st.info(f"Faktor {factor2} tidak memiliki efek yang signifikan (p = {p_values.iloc[1]:.4f} > {alpha}).")
                
                # Jika interaksi signifikan, tampilkan plot interaksi
                if p_values.iloc[2] < alpha:
                    st.write("### Visualisasi Efek Interaksi")
                    
                    # Plot interaksi dengan plotly
                    interaction_data = data.groupby([factor1, factor2], observed=True)[dependent_var].mean().reset_index()
                    
                    fig = px.line(
                        interaction_data, 
                        x=factor1, 
                        y=dependent_var, 
                        color=factor2,
                        markers=True,
                        title=f"Interaksi antara {factor1} dan {factor2}",
                    )
                    
                    fig.update_layout(
                        xaxis_title=factor1,
                        yaxis_title=f"Rata-rata {dependent_var}",
                        legend_title=factor2,
                        height=600
                    )
                    
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Heatmap interaksi
                    pivot_data = interaction_data.pivot(index=factor1, columns=factor2, values=dependent_var)
                    
                    fig = px.imshow(
                        pivot_data,
                        title=f"Heatmap Interaksi {factor1} x {factor2}",
                        labels=dict(x=factor2, y=factor1, color=f"Rata-rata {dependent_var}")
                    )
                    
                    st.plotly_chart(fig, use_container_width=True)
                    
                    st.write("**Menginterpretasikan Hasil Interaksi:**")
                    st.write("""
                        - Jika garis pada plot interaksi sejajar, efek dari satu faktor konsisten di semua level faktor lainnya.
                        - Jika garis bersilangan atau tidak sejajar, ada interaksi yang signifikan. Ini berarti efek satu faktor bergantung pada level faktor lainnya.
                        - Heatmap menunjukkan kombinasi kedua faktor yang menghasilkan nilai tertinggi dan terendah pada variabel dependen.
                    """)
                else:
                    st.info(f"Interaksi antara {factor1} dan {factor2} tidak signifikan (p = {p_values.iloc[2]:.4f} > {alpha}).")
                
                # Kesimpulan
                st.write("## Kesimpulan")
                
                # Buat kesimpulan berdasarkan hasil ANOVA
                conclusions = []
                
                if p_values.iloc[0] < alpha:
                    conclusions.append(f"- **{factor1}** memiliki pengaruh yang signifikan terhadap {dependent_var} (p = {float(p_values.iloc[0]):.4f}, η² = {eta_squared.iloc[0]:.4f}).")
                else:
                    conclusions.append(f"- **{factor1}** tidak memiliki pengaruh yang signifikan terhadap {dependent_var} (p = {float(p_values.iloc[0]):.4f}).")
                
                if p_values.iloc[1] < alpha:
                    conclusions.append(f"- **{factor2}** memiliki pengaruh yang signifikan terhadap {dependent_var} (p = {float(p_values.iloc[1]):.4f}, η² = {eta_squared.iloc[1]:.4f}).")
                else:
                    conclusions.append(f"- **{factor2}** tidak memiliki pengaruh yang signifikan terhadap {dependent_var} (p = {float(p_values.iloc[1]):.4f}).")
                
                if p_values.iloc[2] < alpha:
                    conclusions.append(f"- **Interaksi antara {factor1} dan {factor2}** memiliki pengaruh yang signifikan terhadap {dependent_var} (p = {float(p_values.iloc[2]):.4f}, η² = {eta_squared.iloc[2]:.4f}).")
                else:
                    conclusions.append(f"- **Interaksi antara {factor1} dan {factor2}** tidak memiliki pengaruh yang signifikan terhadap {dependent_var} (p = {float(p_values.iloc[2]):.4f}).")
                
                for conclusion in conclusions:
                    st.markdown(conclusion)
                
                # Tambahkan R-squared
                st.markdown(f"**Model menjelaskan {model.rsquared:.2%} dari variasi dalam {dependent_var}.**")
                
                # Add this after the analysis is complete

                # Create export options
                st.write("## Ekspor Hasil")
                export_col1, export_col2, export_col3 = st.columns(3)

                with export_col1:
                    # CSV export
                    csv = formatted_anova.to_csv(index=True)
                    st.download_button(
                        label="Unduh Tabel ANOVA sebagai CSV",
                        data=csv,
                        file_name=f"anova_results_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                        mime="text/csv"
                    )

                with export_col2:
                    # Excel export option
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                        formatted_anova.to_excel(writer, sheet_name="ANOVA Table", index=True)
                        effects_summary.to_excel(writer, sheet_name="Effect Sizes", index=False)
                        
                        # Add model statistics
                        stats_df = pd.DataFrame({
                            'Metrik': ['R²', 'R² Disesuaikan', 'F-statistic', 'p-value', 'Log-Likelihood', 'AIC', 'BIC'],
                            'Nilai': [
                                f"{model.rsquared:.4f}",
                                f"{model.rsquared_adj:.4f}",
                                f"{model.fvalue:.4f}",
                                f"{model.f_pvalue:.4e}" if model.f_pvalue < 0.0001 else f"{model.f_pvalue:.4f}",
                                f"{model.llf:.4f}",
                                f"{model.aic:.4f}",
                                f"{model.bic:.4f}"
                            ]
                        })
                        stats_df.to_excel(writer, sheet_name="Model Statistics", index=False)
                        
                        # Add summary info
                        summary_info = pd.DataFrame({
                            'Item': ['Dependent Variable', 'Factor 1', 'Factor 2', 'Alpha', 'Sample Size', 'Date'],
                            'Value': [dependent_var, factor1, factor2, str(alpha), str(len(data)), 
                                     datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')]
                        })
                        summary_info.to_excel(writer, sheet_name="Analysis Info", index=False)
                        
                    excel_data = excel_buffer.getvalue()
                    st.download_button(
                        label="Unduh Hasil sebagai Excel",
                        data=excel_data,
                        file_name=f"anova_results_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )

                with export_col3:
                    # Word document export
                    try:
                        docx_buffer = create_docx_report(
                            data, 
                            dependent_var, 
                            factor1, 
                            factor2, 
                            formatted_anova, 
                            effects_summary, 
                            p_values, 
                            eta_squared, 
                            alpha, 
                            model,
                            conclusions
                        )
                        
                        st.download_button(
                            label="Unduh Laporan Lengkap (DOCX)",
                            data=docx_buffer,
                            file_name=f"anova_report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    except Exception as e:
                        st.error(f"Gagal membuat dokumen Word: {str(e)}")
                        st.info("Pastikan library python-docx terinstal: pip install python-docx")

                    # HTML export option (can be opened in Word)
                    try:
                        html_report = create_html_report(
                            dependent_var, 
                            factor1, 
                            factor2, 
                            formatted_anova, 
                            effects_summary, 
                            p_values, 
                            eta_squared, 
                            alpha, 
                            model,
                            conclusions
                        )
                        
                        st.download_button(
                            label="Unduh Laporan Lengkap (HTML)",
                            data=html_report,
                            file_name=f"anova_report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
                            mime="text/html"
                        )
                    except Exception as e:
                        st.error(f"Gagal membuat laporan HTML: {str(e)}")
        else:
            # Memberikan saran untuk pengguna baru
            st.info("Pilih variabel dan klik tombol 'Jalankan Two-Way ANOVA' untuk melihat hasil analisis.")
    
    except Exception as e:
        st.error(f"Terjadi kesalahan: {str(e)}")
        st.write("Pastikan data Anda sesuai untuk analisis ANOVA.")
        st.write("Detail kesalahan untuk membantu pemecahan masalah:")
        st.code(f"{type(e).__name__}: {str(e)}")
        st.write("Tips pemecahan masalah:")
        st.write("1. Pastikan variabel dependen bertipe numerik")
        st.write("2. Pastikan faktor-faktor memiliki setidaknya dua level kategori")
        st.write("3. Pastikan tidak ada nilai yang hilang (NA) di kolom yang dianalisis")
else:
    # Tampilkan instruksi dan contoh saat tidak ada file yang diunggah
    st.info("Silakan unggah file CSV atau Excel yang berisi data untuk analisis Two-Way ANOVA.")
    
    # Menambahkan section contoh data
    with st.expander("Contoh format data yang dibutuhkan"):
        st.write("""
            Untuk melakukan analisis Two-Way ANOVA, Anda memerlukan:
            
            1. Satu variabel dependen (numerik) - misalnya berat badan, produksi susu, pertambahan bobot harian
            2. Dua variabel faktor (kategorikal) - misalnya jenis pakan, breed ternak, sistem pemeliharaan
            
            Contoh format CSV dalam bidang peternakan:
        """)
        
        # Data contoh peternakan yang komprehensif
        example_data = pd.DataFrame({
            'BeratBadan': [
                # Pakan: Konsentrat, Breed: Brahman
                285, 290, 282, 288, 291,
                # Pakan: Konsentrat, Breed: Limousin
                310, 315, 308, 312, 318,
                # Pakan: Konsentrat, Breed: Simental
                325, 330, 328, 332, 335,
                
                # Pakan: Hijauan, Breed: Brahman
                270, 275, 268, 273, 276,
                # Pakan: Hijauan, Breed: Limousin
                295, 298, 293, 297, 301,
                # Pakan: Hijauan, Breed: Simental
                308, 312, 305, 310, 315,
                
                # Pakan: Campuran, Breed: Brahman
                295, 298, 292, 297, 300,
                # Pakan: Campuran, Breed: Limousin
                325, 330, 323, 328, 333,
                # Pakan: Campuran, Breed: Simental
                345, 350, 342, 348, 352
            ],
            'JenisPakan': [
                # 5 Brahman, Konsentrat
                'Konsentrat', 'Konsentrat', 'Konsentrat', 'Konsentrat', 'Konsentrat',
                # 5 Limousin, Konsentrat
                'Konsentrat', 'Konsentrat', 'Konsentrat', 'Konsentrat', 'Konsentrat',
                # 5 Simental, Konsentrat
                'Konsentrat', 'Konsentrat', 'Konsentrat', 'Konsentrat', 'Konsentrat',
                
                # 5 Brahman, Hijauan
                'Hijauan', 'Hijauan', 'Hijauan', 'Hijauan', 'Hijauan',
                # 5 Limousin, Hijauan
                'Hijauan', 'Hijauan', 'Hijauan', 'Hijauan', 'Hijauan',
                # 5 Simental, Hijauan
                'Hijauan', 'Hijauan', 'Hijauan', 'Hijauan', 'Hijauan',
                
                # 5 Brahman, Campuran
                'Campuran', 'Campuran', 'Campuran', 'Campuran', 'Campuran',
                # 5 Limousin, Campuran
                'Campuran', 'Campuran', 'Campuran', 'Campuran', 'Campuran',
                # 5 Simental, Campuran
                'Campuran', 'Campuran', 'Campuran', 'Campuran', 'Campuran'
            ],
            'BreedSapi': [
                # 5 Brahman, Konsentrat
                'Brahman', 'Brahman', 'Brahman', 'Brahman', 'Brahman',
                # 5 Limousin, Konsentrat
                'Limousin', 'Limousin', 'Limousin', 'Limousin', 'Limousin',
                # 5 Simental, Konsentrat
                'Simental', 'Simental', 'Simental', 'Simental', 'Simental',
                
                # 5 Brahman, Hijauan
                'Brahman', 'Brahman', 'Brahman', 'Brahman', 'Brahman',
                # 5 Limousin, Hijauan
                'Limousin', 'Limousin', 'Limousin', 'Limousin', 'Limousin',
                # 5 Simental, Hijauan
                'Simental', 'Simental', 'Simental', 'Simental', 'Simental',
                
                # 5 Brahman, Campuran
                'Brahman', 'Brahman', 'Brahman', 'Brahman', 'Brahman',
                # 5 Limousin, Campuran
                'Limousin', 'Limousin', 'Limousin', 'Limousin', 'Limousin',
                # 5 Simental, Campuran
                'Simental', 'Simental', 'Simental', 'Simental', 'Simental'
            ]
        })
        
        # Tambahkan deskripsi untuk dataset contoh peternakan
        st.markdown("""
        ### Penjelasan Dataset Contoh Peternakan
        
        Dataset ini berisi data berat badan sapi (kg) berdasarkan dua faktor:
        
        1. **JenisPakan**: Jenis pakan yang diberikan pada sapi
           - Konsentrat: Pakan dengan kandungan protein tinggi dan serat rendah
           - Hijauan: Pakan dengan kandungan serat tinggi seperti rumput dan legum
           - Campuran: Kombinasi konsentrat dan hijauan dengan proporsi seimbang
           
        2. **BreedSapi**: Jenis breed (ras) sapi yang dipelihara
           - Brahman: Breed sapi yang beradaptasi baik di daerah tropis dengan kualitas daging medium
           - Limousin: Breed sapi asal Perancis dengan pertumbuhan otot yang baik
           - Simental: Breed sapi asal Swiss dengan produktivitas tinggi
           
        Dataset ini menunjukkan pengaruh jenis pakan, breed sapi, dan interaksi keduanya terhadap berat badan sapi.
        """)
        
        # Tampilkan preview data
        st.write("Preview data contoh:")
        st.write(example_data.head(10))
        
        # Tombol unduh data contoh
        st.download_button(
            label="Unduh Data Contoh",
            data=example_data.to_csv(index=False),
            file_name="contoh_anova_peternakan.csv",
            mime="text/csv"
        )
        
        # Visualisasi data contoh untuk membantu pemahaman
        st.markdown("### Visualisasi Data Contoh")
        
        # Tampilkan boxplot untuk memvisualisasikan perbedaan
        fig = px.box(
            example_data, 
            x="JenisPakan", 
            y="BeratBadan", 
            color="BreedSapi", 
            title="Distribusi Berat Badan Sapi berdasarkan Jenis Pakan dan Breed"
        )
        st.plotly_chart(fig, use_container_width=True)
        
        # Tampilkan plot interaksi
        interaction_data = example_data.groupby(['JenisPakan', 'BreedSapi'])['BeratBadan'].mean().reset_index()
        fig = px.line(
            interaction_data, 
            x="JenisPakan", 
            y="BeratBadan", 
            color="BreedSapi", 
            markers=True,
            title="Plot Interaksi: Pengaruh Jenis Pakan dan Breed terhadap Berat Badan Sapi"
        )
        st.plotly_chart(fig, use_container_width=True)
        
        st.markdown("""
        ### Hasil yang Diharapkan
        
        Dengan dataset ini, Anda dapat mengamati:
        
        1. **Efek utama Jenis Pakan**: Pakan campuran menghasilkan berat badan lebih tinggi dibanding jenis pakan lainnya
        2. **Efek utama Breed Sapi**: Simental memiliki berat badan tertinggi, diikuti Limousin, kemudian Brahman
        3. **Efek interaksi**: Kombinasi pakan campuran dengan breed Simental memberikan hasil terbaik dengan peningkatan yang lebih signifikan dibandingkan kombinasi lainnya
        
        Data ini menunjukkan kondisi riil di lapangan dimana pemilihan jenis pakan yang tepat untuk breed tertentu dapat mengoptimalkan hasil peternakan. Cobalah menganalisis dataset ini untuk melihat signifikansi statistik dari perbedaan yang teramati!
        """)
# Add this import at the top with the other imports
import io

# Sebelum bagian footer (sekitar baris 500)
# Dapatkan tahun saat ini untuk footer
current_year = datetime.datetime.now().year

# Footer with LinkedIn profile link and improved styling
st.markdown("""
<hr style="height:1px;border:none;color:#333;background-color:#333;margin-top:30px;margin-bottom:20px">
""", unsafe_allow_html=True)

st.markdown(f"""
<div style="text-align:center; padding:15px; margin-top:10px; margin-bottom:20px">
    <p style="font-size:16px; color:#555">
        © {current_year} Developed by: 
        <a href="https://www.linkedin.com/in/galuh-adi-insani-1aa0a5105/" target="_blank" 
           style="text-decoration:none; color:#0077B5; font-weight:bold">
            <img src="https://content.linkedin.com/content/dam/me/business/en-us/amp/brand-site/v2/bg/LI-Bug.svg.original.svg" 
                 width="16" height="16" style="vertical-align:middle; margin-right:5px">
            Galuh Adi Insani
        </a> 
        with <span style="color:#e25555">❤️</span>
    </p>
    <p style="font-size:12px; color:#777">All rights reserved.</p>
</div>
""", unsafe_allow_html=True)